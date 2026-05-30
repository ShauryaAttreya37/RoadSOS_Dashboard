"""Build the RoadSoS hackathon Word report."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/sessions/zen-friendly-brown/mnt/outputs"

NAVY = RGBColor(0x0f, 0x34, 0x60)
AMBER = RGBColor(0xf5, 0xa6, 0x23)
RED = RGBColor(0xe6, 0x39, 0x46)
TEAL = RGBColor(0x2a, 0x9d, 0x8f)
GREY = RGBColor(0x47, 0x55, 0x69)
LIGHT_GREY = RGBColor(0x6b, 0x72, 0x80)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    if level == 1:
        r.font.size = Pt(20); r.font.color.rgb = color
    elif level == 2:
        r.font.size = Pt(15); r.font.color.rgb = color
    else:
        r.font.size = Pt(12.5); r.font.color.rgb = color
    return p


def add_para(doc, text, size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullet(doc, text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r2 = p.add_run(" — " + text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def add_image(doc, path, width_inches=6.4, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = LIGHT_GREY
        cap.paragraph_format.space_after = Pt(10)


def build_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        set_cell_shading(cell, '0F3460')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows, start=1):
        row = table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, 'F1F5F9')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_callout_box(doc, title, body, color_hex='FFF3E0'):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, color_hex)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# === Build document ===
doc = Document()

# Page setup
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# === Cover ===
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(80)
r = cover.add_run("RoadSoS")
r.bold = True; r.font.size = Pt(54); r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Smart Helmet SOS + Predictive Safety System")
r.font.size = Pt(20); r.font.color.rgb = AMBER; r.bold = True

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Predicting the crash. Calling for help. Warning the world. Powered by the sun.")
r.font.size = Pt(13); r.italic = True; r.font.color.rgb = GREY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Hackathon Prototype Blueprint  •  2026")
r.font.size = Pt(11); r.font.color.rgb = LIGHT_GREY

# Hero stats table on cover
doc.add_paragraph().paragraph_format.space_before = Pt(40)
stats = doc.add_table(rows=1, cols=4)
stats.autofit = False
stat_data = [("1.5 L", "annual road deaths"), ("44%", "two-wheeler riders"), ("<100 ms", "sense→warn loop"), ("₹4,985", "prototype BOM")]
for i, (val, lbl) in enumerate(stat_data):
    cell = stats.rows[0].cells[i]
    cell.width = Inches(1.6)
    cell.text = ''
    set_cell_shading(cell, '0F3460')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(val)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0xf5, 0xa6, 0x23)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(lbl)
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_page_break()

# === 1. Problem ===
add_heading(doc, "1. The Problem — Why RoadSoS Exists")
add_para(doc, "India loses over 1.5 lakh lives to road accidents every year. Two-wheeler riders account for 44% of all road fatalities (MoRTH, 2023). The single biggest killer is not the crash itself — it is the delay in emergency response.")
add_para(doc, "Studies show that if medical aid reaches the victim within the Golden Hour (first 60 minutes after impact), survival rates jump by over 70%. Today's reality is grim:")
add_bullet(doc, "The rider often loses consciousness, so nobody calls for help.")
add_bullet(doc, "Bystanders don't know the victim's blood group, allergies, or contacts.")
add_bullet(doc, "Following riders on highways collide into the same wreck — secondary collisions.")
add_bullet(doc, "Conventional crash sensors trigger only after impact — too late to prevent it.")
add_para(doc, "RoadSoS flips this. It is a predictive, self-aware, mesh-connected helmet that warns the rider before a crash, calls for help during one, and warns nearby riders after one — entirely on solar-assisted battery power.", italic=True, color=NAVY)

# === 2. Scenario ===
add_heading(doc, "2. Real-Life Scenario — A Day with Ravi")
add_para(doc, "Ravi commutes from Hosur to Bangalore on NH-44. Here is how his RoadSoS helmet works across an ordinary, rainy day.")

scenarios = [
    ("06:30 — Boot & charge.", "Helmet powers on. Green LED pulses — battery 92%, solar trickle active. Medical profile (O+, asthma, contact: wife Anjali) encrypted on ESP32-S3."),
    ("07:12 — Rain & oil patch at km 47.", "IMU detects abnormal lateral acceleration and a 23° lean angle building too fast. The on-device PINN estimates skid probability past 0.65. Buzzer pulses, brow LEDs flash amber, and a haptic motor taps Ravi's temple. He eases the throttle. Crisis averted."),
    ("11:40 — Truck swerves; Ravi is thrown.", "IMU registers >6 g impact + free-fall + sudden tilt >80°. The edge-AI 1D-CNN confirms crash (not pothole, not dropped helmet). A 10-second cancel countdown begins. He doesn't cancel."),
    ("11:40:10 — Emergency mode activates.", "GPS locked. SIM7600 dispatches SMS + HTTPS to 108, hospital API, police, and Anjali. LoRa 433 MHz broadcast — every RoadSoS helmet within ~3 km buzzes. Beacon mode: rear LEDs strobe red, buzzer chirps Morse SOS. A bystander scans the QR on the shell — first-aid + blood group instantly visible."),
    ("11:54 — Ambulance arrives.", "Golden Hour preserved. Three riders 1.5 km behind have already slowed because their helmets vibrated. No secondary collision."),
    ("16:00 — Return ride.", "Daytime sun has been trickle-charging the helmet — battery is at 87% despite 8 hours of operation."),
]
for title, body in scenarios:
    add_callout_box(doc, title, body, 'FFF8EB')

doc.add_page_break()

# === 3. System Architecture ===
add_heading(doc, "3. System Architecture")
add_para(doc, "The helmet is a three-tier stack. The entire safety loop — sense → infer → warn — completes in under 100 milliseconds, well below human reaction time.")
add_image(doc, f"{OUT}/fig_1_architecture.png", caption="Figure 1 — Three-tier system architecture")

# === 4. Sensor Map & BOM ===
add_heading(doc, "4. Sensor Placement & Bill of Materials")
add_para(doc, "Every sensor lives where its physics works best. The IMU sits at the crown for the cleanest gyroscopic reference. The solar film occupies the top curved shell where sky exposure is maximal. The brain (ESP32 + radios + battery) lives in a rear pod, preserving the helmet's center of mass and keeping the Li-Po cell away from the rider's skull.")
add_image(doc, f"{OUT}/fig_2_sensor_map.png", caption="Figure 2 — Sensor placement on the helmet")

add_heading(doc, "4.1 Bill of Materials", level=2)
bom_rows = [
    ["1", "Microcontroller", "ESP32-S3-WROOM-1 (8 MB PSRAM)", "Rear pod, EMI-shielded", "₹450"],
    ["2", "6-axis IMU", "ICM-42688-P (TDK InvenSense)", "Crown, centered", "₹380"],
    ["3", "LoRa radio", "Ra-02 (SX1278, 433 MHz)", "Rear pod + helical antenna", "₹220"],
    ["4", "GNSS", "u-blox NEO-6M", "Top shell, sky-facing", "₹350"],
    ["5", "Cellular modem", "SIM7600E-H 4G LTE", "Rear pod, chin-strap antenna", "₹1,800"],
    ["6", "Solar film", "Flexible a-Si 5V / 200 mA", "Top curved shell", "₹450"],
    ["7", "Battery", "Li-Po 3.7 V 2000 mAh + BMS", "Rear pod, aramid sleeve", "₹350"],
    ["8", "Solar MPPT IC", "CN3791", "Power board", "₹120"],
    ["9", "Buzzer / haptics / LEDs / bone-conduction", "Various", "Temple, ear-cup, brow", "₹400"],
    ["10", "Cancel button + QR + NFC + PCB + misc", "—", "—", "₹465"],
    ["", "TOTAL", "", "", "₹4,985"],
]
build_table(doc, ["#", "Component", "Model", "Placement", "Cost"], bom_rows, col_widths=[0.3, 1.7, 1.8, 1.8, 0.7])
add_para(doc, "Target retail BOM at 10k volume: ₹2,800 → consumer price ₹4,999. Total weight added to a stock helmet: ~120 g.", italic=True, color=LIGHT_GREY, size=10)

doc.add_page_break()

# === 5. AI Pipeline ===
add_heading(doc, "5. Edge-AI Pipeline — Two Models, One Chip")
add_para(doc, "RoadSoS runs two complementary models on the same ESP32-S3, one per FreeRTOS core. The PINN handles the before (prevention). The TinyML classifier handles the during (confirmation). Together they cover the full safety timeline — neither alone would.")
add_image(doc, f"{OUT}/fig_3_ai_pipeline.png", caption="Figure 3 — Edge-AI pipeline with PINN and TinyML co-resident on the ESP32-S3")

add_heading(doc, "5.1 PINN — Physics-Informed Neural Network", level=2)
add_bullet(doc, "predicts rider instability before loss of control. Outputs P(skid), P(tipover), and time-to-critical every 100 ms.", label="What it does")
add_bullet(doc, "a vanilla NN would need millions of labeled crash examples. The PINN bakes the motorcycle dynamic equations into the loss function as soft constraints, so it generalises from small data.", label="Why physics-informed")
add_bullet(doc, "Roll dynamics: I·φ̈ = m·g·h·sin(φ) − m·v²·h·cos(φ)·δ/L (lean φ, steering δ, wheelbase L); Friction circle: √(a_x² + a_y²) ≤ μ·g; simplified Pacejka tire saturation curve.", label="Physics encoded")
add_bullet(doc, "MLP, 4 hidden layers × 32 neurons; 12 KB INT8; ESP32-S3 core 1; ~3 ms inference.", label="Where it lives")

add_heading(doc, "5.2 TinyML — 1D-CNN Crash Classifier", level=2)
add_bullet(doc, "confirms whether an impact event is a real crash vs. dropped helmet, pothole, speed bump, or hand-slap.", label="What it does")
add_bullet(doc, "1D-CNN with 3 conv + 2 dense layers, INT8-quantized. Input: 200 ms sliding window of 6-axis IMU (60 × 6 tensor). Output: softmax over {crash, fall, bump, normal}.", label="Architecture")
add_bullet(doc, "Public IMU datasets (SisFall, UMA Fall) + 1,200 self-collected motorcycle runs over rough Indian roads.", label="Training data")
add_bullet(doc, "48 KB INT8; ESP32-S3 core 0 with ESP-NN acceleration via TensorFlow Lite Micro; ~9 ms inference.", label="Where it lives")

add_callout_box(doc, "Why both, not one?", "If only the PINN ran, we would get too many false alarms (no impact confirmation). If only the CNN ran, we would arrive too late (post-impact only). The two together cover the full timeline.", "FFF3E0")

doc.add_page_break()

# === 6. Emergency flow ===
add_heading(doc, "6. Emergency Response Flow")
add_para(doc, "From impact to ambulance dispatch in under 15 seconds. Every millisecond is accounted for.")
add_image(doc, f"{OUT}/fig_4_emergency_flow.png", caption="Figure 4 — Emergency response timeline")

steps = [
    ("t = 0 ms — IMPACT", "IMU registers >6 g, free-fall flag, tilt >80°."),
    ("t = 9 ms — CONFIRM", "TinyML 1D-CNN softmax votes 'crash' with confidence >0.85."),
    ("t = 12 ms — PROMPT", "Buzzer pulses 'Cancel?', temple haptic taps the rider."),
    ("t = 0 – 10 s — WINDOW", "Rider can press chin-strap cancel button. GPS warm-starts in parallel."),
    ("t = 11 s — DISPATCH", "SOS packet assembled: {lat, lon, time, rider_id, blood_group, contacts, impact_g}."),
    ("t = 11 s — FOUR CHANNELS IN PARALLEL", "Cellular (108 SMS + hospital API HTTPS POST + contact SMS) · LoRa mesh broadcast · BLE phone redundancy · On-helmet beacon (LED strobe + buzzer Morse + NFC active)."),
    ("t = ongoing — FORENSIC LOG", "30 s of pre-impact telemetry stored on PSRAM for insurance / law-enforcement playback."),
]
for t, b in steps:
    add_bullet(doc, b, label=t)

doc.add_page_break()

# === 7. LoRa Mesh ===
add_heading(doc, "7. LoRa Rider-to-Rider Safety Mesh")
add_para(doc, "When one helmet crashes, every helmet within ~3 km knows in under 2 seconds. No cell tower required — works in dead zones and rural areas.")
add_image(doc, f"{OUT}/fig_5_lora_mesh.png", caption="Figure 5 — LoRa rider-to-rider safety mesh")
add_bullet(doc, "433 MHz (WPC license-free in India).", label="Frequency")
add_bullet(doc, "~3 km line-of-sight, ~800 m urban.", label="Range")
add_bullet(doc, "lightweight broadcast — every helmet listens on a shared channel; received packets are re-broadcast once (TTL = 2) to extend range without flooding.", label="Protocol")
add_bullet(doc, "32 bytes — type, TTL, lat, lon, severity, rider ID, hash.", label="Packet payload")
add_bullet(doc, "their helmet vibrates, brow LEDs flash amber, and (if BLE-paired phone is mounted) a Google Maps pin appears showing the hazard location.", label="What nearby riders see")
add_bullet(doc, "on highways, secondary collisions kill nearly as many people as the primary crash. A 5-second early warning is enough to brake from 80 km/h to a survivable speed.", label="Why this matters")

doc.add_page_break()

# === 8. Solar Charging ===
add_heading(doc, "8. Solar Charging Subsystem — Powered by the Sun")
add_para(doc, "A 2,000 mAh battery gives ~18 hours of active monitoring. Riders often forget to charge. Solar trickle-charging during daytime rides extends usable life to effectively infinite for daily commutes.")
add_image(doc, f"{OUT}/fig_6_solar_charging.png", caption="Figure 6 — Solar charging subsystem and energy budget")

add_heading(doc, "8.1 How it works", level=2)
add_bullet(doc, "(5V, 200 mA peak) laminated to the top curved surface using transparent automotive-grade adhesive.", label="Solar film")
add_bullet(doc, "feeds an MPPT controller (CN3791) which extracts maximum power even under partial shading from the visor or head tilt.", label="MPPT")
add_bullet(doc, "MPPT output charges the 3.7 V Li-Po cell via a TP4056-class linear charger with over-voltage / over-current / over-temperature cutoffs.", label="Charger IC")
add_bullet(doc, "informs the ESP32 whether it's day or night — at night the helmet aggressively duty-cycles the GPS to save power; during the day it can keep all sensors hot.", label="Lux sensor")

add_heading(doc, "8.2 Energy budget", level=2)
energy_rows = [
    ["Active (all sensors, AI running)", "~110 mA", "~18 h", "+ ~200 mA in sun → net positive"],
    ["Idle (Wake-on-Motion only)", "~12 mA", "~160 h", "Recharges in <2 h sun"],
    ["Emergency (LoRa TX + cellular)", "~480 mA burst", "~30 min sustained", "Battery reserve only"],
    ["6 h daytime ride (active + sun)", "−110 mA, +200 mA", "Battery GAINS ~540 mAh", "Effectively unlimited"],
]
build_table(doc, ["Mode", "Current", "Battery life (2000 mAh)", "Solar offset"], energy_rows, col_widths=[2.0, 1.2, 1.7, 1.7])

add_heading(doc, "8.3 Safety", level=2)
add_bullet(doc, "Battery sits in a fireproof aramid sleeve at the rear pod (away from rider's skull).")
add_bullet(doc, "BMS has thermal cutoff at 55 °C.")
add_bullet(doc, "Solar panel is a non-shattering thin film — no glass anywhere on the helmet.")

doc.add_page_break()

# === 9. Why it wins / roadmap ===
add_heading(doc, "9. Why RoadSoS Wins — Judging Criteria Map")
judging_rows = [
    ["Innovation", "First helmet to combine PINN + TinyML on a sub-₹500 MCU. Solar trickle is unique in this category."],
    ["Feasibility", "All components off-the-shelf. Working prototype achievable in 6 weeks. Total BOM under ₹5,000."],
    ["Impact", "Targets 44% of Indian road fatalities. Even a 10% Golden Hour improvement saves ~6,500 lives/year."],
    ["Scalability", "LoRa mesh gets more valuable as more helmets deploy — classic network effect."],
    ["Sustainability", "Solar reduces charging burden, extends battery life, lowers e-waste."],
    ["UX", "Zero rider input required. Helmet just works from the moment it's worn."],
]
build_table(doc, ["Criterion", "How RoadSoS scores"], judging_rows, col_widths=[1.5, 5.0])

add_heading(doc, "10. Roadmap")
roadmap = [
    ("Weeks 0–2 (Hackathon)", "Working prototype on one helmet — IMU + ESP32 + LoRa + buzzer + GPS + cellular dummy."),
    ("Weeks 3–6", "Train PINN on real motorcycle data; collaborate with local riding clubs."),
    ("Month 3", "Pilot with 50 delivery riders (Dunzo / Swiggy)."),
    ("Month 6", "Integrate with 108 EMS official API + Karnataka highway police."),
    ("Year 1", "B2B sales to logistics fleets; B2C launch at ₹4,999."),
    ("Year 2", "Expand to construction-worker helmets, with the same predictive-safety + SOS stack."),
]
for t, b in roadmap:
    add_bullet(doc, b, label=t)

add_heading(doc, "11. The Tagline", color=AMBER)
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run("RoadSoS — Predicting the crash. Calling for help. Warning the world. Powered by the sun.")
r.bold = True
r.font.size = Pt(15)
r.italic = True
r.font.color.rgb = NAVY

# Save
out_path = f"{OUT}/RoadSoS_Report.docx"
doc.save(out_path)
print("Wrote:", out_path)
