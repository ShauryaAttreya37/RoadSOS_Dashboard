from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from roadsos_app.modules.emergency_numbers import EMERGENCY_NUMBERS
SERVICE_QUERIES = {
    "hospital": 'node["amenity"="hospital"]',
    "clinic": 'node["amenity"="clinic"]',
    "police": 'node["amenity"="police"]',
    "fire_station": 'node["amenity"="fire_station"]',
    "ambulance": 'node["emergency"="ambulance_station"]',
    "towing": 'node["amenity"="vehicle_inspection"]',
    "puncture_shop": 'node["shop"="tyres"]',
    "car_showroom": 'node["shop"="car"]',
}


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "SmartHelmetSOS_Documentation.docx"


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 13)
    return paragraph


def add_para(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(f"- {text}")
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
    return table


def main():
    doc = Document(DOCX)
    remove_existing_addendum(doc)
    doc.add_page_break()

    title = doc.add_paragraph("RoadSoS Streamlit App Addendum")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(20)
    add_para(
        doc,
        "This addendum documents the production Streamlit app built under roadsos_app, including the new global location and emergency services layer.",
    )

    add_heading(doc, "1. System Overview")
    add_para(
        doc,
        "RoadSoS is organized as a two-layer system. The helmet intelligence layer simulates helmet IMU streams, predicts skid risk through a Physics-Informed Neural Network, and estimates injury risk through HIC15 and BrIC. The location services layer takes GPS coordinates, fetches nearby emergency and vehicle-support contacts from OpenStreetMap, and stores a 24-hour offline cache for fallback use.",
    )

    add_heading(doc, "2. Module Reference")
    rows = [
        ("app.py", "Streamlit entry point", "Navigation shell and project overview", "Streamlit page"),
        ("pages/1_PINN_Dashboard.py", "Helmet AI dashboard", "Scenario, duration, rider profile", "Dark-mode plots, status banner, metrics"),
        ("pages/2_Nearby_Services.py", "Location services finder", "Country, radius, latitude, longitude", "Emergency numbers, service tables, cache status"),
        ("pages/3_Emergency_Profile.py", "Rider medical profile", "Medical fields and last known location", "SOS JSON packet and QR code"),
        ("modules/simulation.py", "Synthetic IMU generator", "Scenario parameters", "IMU arrays and physics targets"),
        ("modules/pinn_model.py", "PINN model and normaliser", "Normalized IMU tensor", "Physical predictions"),
        ("modules/injury_metrics.py", "Injury and skid metrics", "Acceleration, gyro, friction", "HIC15, BrIC, severity, P(skid)"),
        ("modules/losses.py", "Physics residual losses", "Model outputs and tensors", "Data, vehicle, biomechanics losses"),
        ("modules/train.py", "Training routine", "Dataset arrays", "Checkpoint and loss history"),
        ("modules/inference.py", "Cached model loading and inference", "Simulation function and kwargs", "Dashboard-ready outputs"),
        ("modules/nearby_services.py", "Overpass API client", "GPS coordinates and radius", "Categorized nearby contacts"),
        ("modules/offline_cache.py", "Offline cache", "Cache key and data", "Fresh cached JSON payloads"),
        ("modules/emergency_numbers.py", "Country emergency numbers", "ISO country code", "Ambulance, police, fire, unified numbers"),
    ]
    add_table(doc, ["File", "Purpose", "Inputs", "Outputs"], rows)

    add_heading(doc, "3. PINN Architecture")
    add_para(
        doc,
        "The PINN receives seven inputs: time plus three accelerometer and three gyroscope axes. It predicts velocity, lean angle, effective friction, and three brain displacement states. The fully connected tanh network uses Xavier initialization. Physics residuals constrain vehicle friction dynamics and Kelvin-Voigt brain biomechanics.",
    )
    add_para(doc, "Loss = data MSE + 0.1 vehicle residual loss + 0.05 biomechanics residual loss. Training uses Adam, cosine annealing, and gradient clipping at 1.0.")

    add_heading(doc, "4. Injury Metrics")
    add_para(doc, "HIC15 is computed over the highest-risk 15 ms resultant acceleration window: HIC = Δt × mean(a)^2.5, with acceleration expressed in g.")
    add_para(doc, "BrIC combines peak angular velocities around the x, y, and z axes against NHTSA critical thresholds of 66.3, 56.5, and 42.2 rad/s.")
    add_para(doc, "Severity is LOW below HIC 700 and BrIC 0.5, MODERATE at elevated risk, and SEVERE above HIC 1000 or BrIC 1.0.")

    add_heading(doc, "5. Location Services")
    add_para(
        doc,
        "The services layer builds one combined Overpass QL query using [out:json][timeout:25] and around-radius filters for each supported service type. Results are parsed into normalized records containing name, latitude, longitude, distance_km, phone, and address.",
    )
    add_table(doc, ["Category", "OSM selector"], sorted(SERVICE_QUERIES.items()))
    add_para(
        doc,
        "Caching is performed before live fetch. The cache key is lat_lon_radius rounded to three decimal places for coordinates. Fresh cache entries are reused for 24 hours and stored under roadsos_app/data/cache.",
    )

    add_heading(doc, "6. Emergency Numbers")
    add_table(
        doc,
        ["Country", "Ambulance", "Police", "Fire", "Unified"],
        [
            (country, values["ambulance"], values["police"], values["fire"], values["unified"])
            for country, values in sorted(EMERGENCY_NUMBERS.items())
        ],
    )

    add_heading(doc, "7. Evaluation Criteria Coverage")
    criteria = [
        ("Reliability and accuracy", "OSM-backed services plus cache freshness metadata on the services page."),
        ("Number of contacts", "The services page displays a total count of contacts found within the selected radius."),
        ("Offline functionality", "offline_cache.py persists 24-hour JSON caches for reuse without another live API call."),
        ("Innovation", "PINN prediction combines data loss with vehicle and biomechanics physics residuals."),
        ("Global applicability", "Country selector and international emergency number table avoid India-only logic."),
    ]
    add_table(doc, ["Criterion", "RoadSoS coverage"], criteria)

    add_heading(doc, "8. How to Run")
    add_para(doc, 'From the project root: pip install -r roadsos_app/requirements.txt')
    add_para(doc, 'Launch the dashboard: streamlit run roadsos_app/app.py')
    add_para(doc, "Use the sidebar to open Helmet AI Dashboard, Nearby Emergency Services, or Rider Emergency Profile.")

    add_heading(doc, "9. Hardware BOM")
    for item in [
        "ESP32-S3 edge controller for TinyML/PINN inference",
        "ICM-42688-P high-range IMU",
        "Ra-02 LoRa radio for nearby-rider warning mesh",
        "SIM7600 cellular module for emergency alert uplink",
        "NEO-6M GPS module for location fix",
        "Flexible solar film for trickle charging",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Deployment Roadmap")
    add_table(
        doc,
        ["Phase", "Focus"],
        [
            ("Week 0-2 hackathon", "Complete dashboard, demo flows, cache-enabled emergency services, and pitch artifacts."),
            ("Month 3 pilot", "Collect road data with riders, validate service lookup accuracy, and tune crash/skid thresholds."),
            ("Year 1 B2B", "Fleet integrations, responder workflows, ruggedized hardware, and OTA model updates."),
        ],
    )

    doc.save(DOCX)


def remove_existing_addendum(doc):
    body = doc.element.body
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t"))
        if "RoadSoS Streamlit App Addendum" in text:
            start = idx
            break
    if start is None:
        return
    for child in children[start:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


if __name__ == "__main__":
    main()
